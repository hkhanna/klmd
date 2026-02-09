"""
Tests for KLMD Docx Renderer
"""

import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document

from klmd.parser import (
    CommentNode,
    CrossReferenceNode,
    DefinedTermNode,
    DocumentNode,
    ParagraphNode,
    SectionNode,
    SignatureBlockNode,
    TextNode,
    TitleNode,
)
from klmd.renderers.docx import DocxRenderer
from klmd.renderers.docx_config import DocxConfig, StyleMapping
from klmd.renderers.markdown import NumberingScheme


class TestStyleMapping:
    """Tests for StyleMapping."""

    def test_default_styles(self) -> None:
        """Test default style values."""
        mapping = StyleMapping()
        assert mapping.document_title == "Title"
        assert mapping.paragraph == "Normal"
        assert mapping.section_level_1 == "Heading 1"

    def test_get_section_style_valid_levels(self) -> None:
        """Test getting styles for valid levels 1-5."""
        mapping = StyleMapping()
        assert mapping.get_section_style(1) == "Heading 1"
        assert mapping.get_section_style(2) == "Heading 2"
        assert mapping.get_section_style(3) == "Heading 3"
        assert mapping.get_section_style(4) == "Heading 4"
        assert mapping.get_section_style(5) == "Heading 5"

    def test_get_section_style_clamping(self) -> None:
        """Test that levels are clamped to 1-5."""
        mapping = StyleMapping()
        # Level 0 or negative should clamp to 1
        assert mapping.get_section_style(0) == "Heading 1"
        assert mapping.get_section_style(-1) == "Heading 1"
        # Level > 5 should clamp to 5
        assert mapping.get_section_style(6) == "Heading 5"
        assert mapping.get_section_style(100) == "Heading 5"


class TestDocxConfig:
    """Tests for DocxConfig."""

    def test_default_config(self) -> None:
        """Test default configuration values."""
        config = DocxConfig()
        assert config.template_path is None
        assert config.generate_bookmarks is True
        assert config.generate_hyperlinks is True
        assert config.include_comments is False
        assert config.defined_term_bold is True
        assert config.uppercase_entity_names is True
        assert config.cross_ref_template == "Section {number}"

    def test_custom_config(self) -> None:
        """Test custom configuration."""
        config = DocxConfig(
            include_comments=True,
            defined_term_bold=False,
            cross_ref_template="See {number}",
        )
        assert config.include_comments is True
        assert config.defined_term_bold is False
        assert config.cross_ref_template == "See {number}"


class TestDocxRendererBasic:
    """Basic tests for DocxRenderer."""

    def test_render_returns_bytes(self) -> None:
        """Test that render returns bytes."""
        document = DocumentNode(children=[])
        renderer = DocxRenderer()
        result = renderer.render(document)
        assert isinstance(result, bytes)

    def test_render_produces_valid_docx(self) -> None:
        """Test that output is a valid docx (ZIP) file."""
        document = DocumentNode(children=[])
        renderer = DocxRenderer()
        result = renderer.render(document)

        # Check ZIP magic bytes (docx is a ZIP file)
        assert result[:4] == b"PK\x03\x04"

        # Verify it can be opened as a ZIP
        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            assert "word/document.xml" in zf.namelist()

    def test_render_parseable_by_python_docx(self) -> None:
        """Test that output can be parsed by python-docx."""
        document = DocumentNode(children=[])
        renderer = DocxRenderer()
        result = renderer.render(document)

        buffer = BytesIO(result)
        doc = Document(buffer)
        assert doc is not None


class TestDocxRendererTitles:
    """Tests for title rendering."""

    def test_document_title(self) -> None:
        """Test rendering a document title."""
        title = TitleNode(
            title="Test Document",
            is_document_title=True,
            has_attachment_placeholder=False,
            subtitle=None,
            children=[],
        )
        document = DocumentNode(children=[title])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Test Document"

    def test_attachment_title_with_number(self) -> None:
        """Test rendering an attachment title with number."""
        title = TitleNode(
            title="Exhibit",
            is_document_title=False,
            has_attachment_placeholder=True,
            subtitle=None,
            children=[],
        )
        document = DocumentNode(children=[title])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        # Default attachment numbering is "letters" which gives "A"
        assert doc.paragraphs[0].text == "Exhibit A"

    def test_attachment_title_with_subtitle(self) -> None:
        """Test rendering an attachment title with subtitle."""
        title = TitleNode(
            title="Exhibit",
            is_document_title=False,
            has_attachment_placeholder=True,
            subtitle="Pricing Terms",
            children=[],
        )
        document = DocumentNode(children=[title])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "Exhibit A"
        assert doc.paragraphs[1].text == "Pricing Terms"


class TestDocxRendererSections:
    """Tests for section rendering."""

    def test_section_with_title(self) -> None:
        """Test rendering a section with title."""
        section = SectionNode(
            level=1,
            title="Introduction",
            children=[],
        )
        document = DocumentNode(children=[section])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        # Default decimal numbering gives "1."
        assert doc.paragraphs[0].text == "1. Introduction"

    def test_section_without_title(self) -> None:
        """Test rendering a section without title."""
        section = SectionNode(
            level=1,
            title=None,
            children=[TextNode(text="Some content here.")],
        )
        document = DocumentNode(children=[section])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "1."
        assert doc.paragraphs[1].text == "Some content here."

    def test_nested_sections(self) -> None:
        """Test rendering nested sections."""
        section1 = SectionNode(level=1, title="First", children=[])
        section2 = SectionNode(level=2, title="Nested", children=[])
        document = DocumentNode(children=[section1, section2])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 2
        assert doc.paragraphs[0].text == "1. First"
        assert doc.paragraphs[1].text == "1.1. Nested"

    def test_section_numbering_presets(self) -> None:
        """Test different numbering presets."""
        section = SectionNode(level=1, title="Test", children=[])
        document = DocumentNode(children=[section])

        # Test legal preset
        config = DocxConfig(
            section_numbering=NumberingScheme.from_preset("legal")
        )
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        # Legal preset uses arabic without suffix at level 1
        assert doc.paragraphs[0].text == "1. Test"


class TestDocxRendererParagraphs:
    """Tests for paragraph rendering."""

    def test_simple_paragraph(self) -> None:
        """Test rendering a simple paragraph."""
        para = ParagraphNode(children=[TextNode(text="Hello, world!")])
        document = DocumentNode(children=[para])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Hello, world!"

    def test_paragraph_with_defined_term(self) -> None:
        """Test rendering a paragraph with a defined term."""
        para = ParagraphNode(children=[
            TextNode(text="This Agreement "),
            DefinedTermNode(term="Agreement", descriptor="the"),
            TextNode(text=" is binding."),
        ])
        document = DocumentNode(children=[para])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "This Agreement (the Agreement) is binding."

    def test_defined_term_bold(self) -> None:
        """Test that defined terms are bold by default."""
        para = ParagraphNode(children=[
            DefinedTermNode(term="Term", descriptor=None),
        ])
        document = DocumentNode(children=[para])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        # Check that the term run is bold
        runs = doc.paragraphs[0].runs
        # Find the run containing "Term"
        term_run = None
        for run in runs:
            if "Term" in run.text:
                term_run = run
                break
        assert term_run is not None
        assert term_run.bold is True

    def test_defined_term_not_bold(self) -> None:
        """Test that defined terms are not bold when configured."""
        para = ParagraphNode(children=[
            DefinedTermNode(term="Term", descriptor=None),
        ])
        document = DocumentNode(children=[para])

        config = DocxConfig(defined_term_bold=False)
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        runs = doc.paragraphs[0].runs
        term_run = None
        for run in runs:
            if "Term" in run.text:
                term_run = run
                break
        assert term_run is not None
        assert term_run.bold is not True


class TestDocxRendererCrossReferences:
    """Tests for cross-reference rendering."""

    def test_cross_reference_resolved(self) -> None:
        """Test rendering a resolved cross-reference."""
        section = SectionNode(level=1, title="Definitions", children=[])
        para = ParagraphNode(children=[
            TextNode(text="See "),
            CrossReferenceNode(
                reference_key="definitions", original_text="[#Definitions]"
            ),
            TextNode(text=" for details."),
        ])
        document = DocumentNode(children=[section, para])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 2
        # The cross-reference should resolve to "Section 1"
        assert "Section 1" in doc.paragraphs[1].text

    def test_cross_reference_custom_template(self) -> None:
        """Test cross-reference with custom template."""
        section = SectionNode(level=1, title="Definitions", children=[])
        para = ParagraphNode(children=[
            CrossReferenceNode(
                reference_key="definitions", original_text="[#Definitions]"
            ),
        ])
        document = DocumentNode(children=[section, para])

        config = DocxConfig(cross_ref_template="See clause {number}")
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert "See clause 1" in doc.paragraphs[1].text

    def test_cross_reference_unresolved(self) -> None:
        """Test that unresolved cross-references render original text."""
        para = ParagraphNode(children=[
            CrossReferenceNode(reference_key="missing", original_text="[#Missing]"),
        ])
        document = DocumentNode(children=[para])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert doc.paragraphs[0].text == "[#Missing]"


class TestDocxRendererComments:
    """Tests for comment rendering."""

    def test_comments_excluded_by_default(self) -> None:
        """Test that comments are excluded by default."""
        comment = CommentNode(content="This is a comment", is_inline=False)
        document = DocumentNode(children=[comment])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        # No paragraphs should be rendered
        assert len(doc.paragraphs) == 0

    def test_comments_included_when_configured(self) -> None:
        """Test that comments are included when configured."""
        comment = CommentNode(content="This is a comment", is_inline=False)
        document = DocumentNode(children=[comment])

        config = DocxConfig(include_comments=True)
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "This is a comment"

    def test_inline_comments_excluded(self) -> None:
        """Test that inline comments are excluded by default."""
        para = ParagraphNode(children=[
            TextNode(text="Some text "),
            CommentNode(content="inline comment", is_inline=True),
            TextNode(text=" more text."),
        ])
        document = DocumentNode(children=[para])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert doc.paragraphs[0].text == "Some text  more text."

    def test_inline_comments_included(self) -> None:
        """Test that inline comments are included when configured."""
        para = ParagraphNode(children=[
            TextNode(text="Some text "),
            CommentNode(content="inline comment", is_inline=True),
            TextNode(text=" more text."),
        ])
        document = DocumentNode(children=[para])

        config = DocxConfig(include_comments=True)
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        assert "[inline comment]" in doc.paragraphs[0].text


class TestDocxRendererSignatureBlocks:
    """Tests for signature block rendering."""

    def test_individual_signature(self) -> None:
        """Test rendering an individual signature block."""
        sig = SignatureBlockNode(
            party_name="John Smith",
            is_entity=False,
            by_entities=[],
            signatory=None,
            fields={"Address": "123 Main St"},
        )
        document = DocumentNode(children=[sig])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        # Should have signature line, name, and address field
        texts = [p.text for p in doc.paragraphs]
        assert any("_" in t for t in texts)  # Signature line
        assert "John Smith" in texts
        assert "Address: 123 Main St" in texts

    def test_entity_signature(self) -> None:
        """Test rendering an entity signature block."""
        sig = SignatureBlockNode(
            party_name="Acme Corp",
            is_entity=True,
            by_entities=[],
            signatory="Jane Doe",
            fields={"Title": "CEO"},
        )
        document = DocumentNode(children=[sig])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        # Entity name should be uppercase by default
        assert "ACME CORP" in texts
        assert "Name: Jane Doe" in texts
        assert "Title: CEO" in texts

    def test_entity_signature_lowercase(self) -> None:
        """Test entity signature without uppercase."""
        sig = SignatureBlockNode(
            party_name="Acme Corp",
            is_entity=True,
            by_entities=[],
            signatory="Jane Doe",
            fields={},
        )
        document = DocumentNode(children=[sig])

        config = DocxConfig(uppercase_entity_names=False)
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert "Acme Corp" in texts
        assert "ACME CORP" not in texts

    def test_entity_signature_with_by_entities(self) -> None:
        """Test entity signature with nested entity chain."""
        sig = SignatureBlockNode(
            party_name="Parent Corp",
            is_entity=True,
            by_entities=["Subsidiary LLC", "Division Inc"],
            signatory="John Manager",
            fields={},
        )
        document = DocumentNode(children=[sig])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert "By: Subsidiary LLC" in texts
        assert "By: Division Inc" in texts


class TestDocxRendererBookmarks:
    """Tests for bookmark generation."""

    def test_bookmark_name_generation(self) -> None:
        """Test bookmark name generation from titles."""
        renderer = DocxRenderer()

        # Simple case
        assert renderer._generate_bookmark_name("Introduction") == "introduction"

        # With spaces
        result = renderer._generate_bookmark_name("Terms and Conditions")
        assert result == "terms_and_conditions"

        # With special characters
        assert renderer._generate_bookmark_name("Article (1)") == "article_1"

        # Starting with number - should add prefix
        assert renderer._generate_bookmark_name("1st Section") == "ref_1st_section"

        # Empty string - should add prefix
        assert renderer._generate_bookmark_name("") == "ref_"

    def test_bookmark_name_truncation(self) -> None:
        """Test that bookmark names are truncated to 40 chars."""
        renderer = DocxRenderer()
        long_title = "A" * 100
        name = renderer._generate_bookmark_name(long_title)
        assert len(name) <= 40

    def test_bookmarks_can_be_disabled(self) -> None:
        """Test that bookmarks can be disabled."""
        title = TitleNode(
            title="Test",
            is_document_title=True,
            has_attachment_placeholder=False,
            subtitle=None,
            children=[],
        )
        document = DocumentNode(children=[title])

        config = DocxConfig(generate_bookmarks=False)
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        # Parse the document XML to check for bookmarks
        buffer = BytesIO(result)
        with zipfile.ZipFile(buffer, "r") as zf:
            doc_xml = zf.read("word/document.xml")
            assert b"bookmarkStart" not in doc_xml


class TestDocxRendererIntegration:
    """Integration tests with full documents."""

    def test_full_document(self) -> None:
        """Test rendering a complete document."""
        doc_title = TitleNode(
            title="Service Agreement",
            is_document_title=True,
            has_attachment_placeholder=False,
            subtitle=None,
            children=[],
        )
        section1 = SectionNode(
            level=1,
            title="Definitions",
            children=[
                ParagraphNode(children=[
                    TextNode(text="The following terms "),
                    DefinedTermNode(term="Agreement", descriptor="this"),
                    TextNode(text=" shall have the meanings set forth below."),
                ]),
            ],
        )
        section2 = SectionNode(
            level=1,
            title="Services",
            children=[
                ParagraphNode(children=[
                    TextNode(text="Subject to the terms of "),
                    CrossReferenceNode(
                        reference_key="definitions", original_text="[#Definitions]"
                    ),
                    TextNode(text="."),
                ]),
            ],
        )
        attachment = TitleNode(
            title="Exhibit",
            is_document_title=False,
            has_attachment_placeholder=True,
            subtitle="Service Levels",
            children=[],
        )
        document = DocumentNode(children=[doc_title, section1, section2, attachment])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))

        # Verify document structure
        texts = [p.text for p in doc.paragraphs]
        assert "Service Agreement" in texts
        assert any("1. Definitions" in t for t in texts)
        assert any("2. Services" in t for t in texts)
        assert "Exhibit A" in texts
        assert "Service Levels" in texts
        # Check cross-reference resolved
        assert any("Section 1" in t for t in texts)

    def test_multiple_attachments(self) -> None:
        """Test rendering multiple attachments."""
        att1 = TitleNode(
            title="Exhibit",
            is_document_title=False,
            has_attachment_placeholder=True,
            subtitle="First",
            children=[],
        )
        att2 = TitleNode(
            title="Exhibit",
            is_document_title=False,
            has_attachment_placeholder=True,
            subtitle="Second",
            children=[],
        )
        document = DocumentNode(children=[att1, att2])

        renderer = DocxRenderer()
        result = renderer.render(document)

        doc = Document(BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert "Exhibit A" in texts
        assert "Exhibit B" in texts
        assert "First" in texts
        assert "Second" in texts


class TestDocxRendererTemplate:
    """Tests for template loading."""

    def test_template_loading(self, tmp_path: Path) -> None:
        """Test that templates can be loaded."""

        # Create a template document
        template = Document()
        template.save(str(tmp_path / "template.docx"))

        # Use it as template
        title = TitleNode(
            title="Test",
            is_document_title=True,
            has_attachment_placeholder=False,
            subtitle=None,
            children=[],
        )
        document = DocumentNode(children=[title])

        config = DocxConfig(template_path=tmp_path / "template.docx")
        renderer = DocxRenderer(config)
        result = renderer.render(document)

        # Should produce valid docx
        doc = Document(BytesIO(result))
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "Test"
