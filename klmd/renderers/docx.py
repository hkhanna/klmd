"""
KLMD Docx Renderer

Converts KLMD AST to Word docx format with configurable styling.
Uses python-docx for document generation and supports Word templates.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from ..parser import (
    CommentNode,
    CrossReferenceNode,
    DefinedTermNode,
    DocumentNode,
    Node,
    ParagraphNode,
    SectionNode,
    SignatureBlockNode,
    TextNode,
    TitleNode,
)
from .docx_config import DocxConfig
from .markdown import MarkdownConfig, NumberingResolver

if TYPE_CHECKING:
    from docx.document import Document as DocumentType


class DocxRenderer:
    """Main docx renderer class."""

    def __init__(self, config: DocxConfig | None = None) -> None:
        self.config = config or DocxConfig()
        self._document: DocumentType | None = None
        self._resolver: NumberingResolver | None = None

    def render(self, document: DocumentNode) -> bytes:
        """Render document to docx bytes."""
        # Phase 1: Resolve numbering (reuse NumberingResolver from markdown)
        md_config = MarkdownConfig(
            section_numbering=self.config.section_numbering,
            attachment_numbering=self.config.attachment_numbering,
        )
        self._resolver = NumberingResolver(md_config)
        self._resolver.resolve(document)

        # Load template or create blank document
        if self.config.template_path:
            self._document = Document(str(self.config.template_path))
        else:
            self._document = Document()

        # Phase 2: Render nodes
        for child in document.children:
            self._render_node(child)

        # Return as bytes
        buffer = BytesIO()
        self._document.save(buffer)
        return buffer.getvalue()

    def _render_node(self, node: Node) -> None:
        """Render a single node to the document."""
        if isinstance(node, TitleNode):
            self._render_title(node)
        elif isinstance(node, SectionNode):
            self._render_section(node)
        elif isinstance(node, ParagraphNode):
            self._render_paragraph(node)
        elif isinstance(node, SignatureBlockNode):
            self._render_signature_block(node)
        elif isinstance(node, CommentNode):
            self._render_comment(node)
        # TextNode, DefinedTermNode, CrossReferenceNode are handled inline

    def _render_title(self, node: TitleNode) -> None:
        """Render a title node."""
        assert self._document is not None
        assert self._resolver is not None

        styles = self.config.paragraph_styles

        if node.has_attachment_placeholder:
            # Attachment title with number
            number = self._resolver.attachment_numbers[id(node)]
            title_text = f"{node.title} {number}"

            paragraph = self._document.add_paragraph(style=styles.attachment_title)
            paragraph.add_run(title_text)

            # Add bookmark for cross-references
            if self.config.generate_bookmarks:
                self._add_bookmark_to_paragraph(paragraph, node.title)

            # Add subtitle if present
            if node.subtitle:
                subtitle_para = self._document.add_paragraph(
                    style=styles.attachment_subtitle
                )
                subtitle_para.add_run(node.subtitle)

                # Also add bookmark for subtitle
                if self.config.generate_bookmarks:
                    self._add_bookmark_to_paragraph(subtitle_para, node.subtitle)
        else:
            # Regular document title
            paragraph = self._document.add_paragraph(style=styles.document_title)
            paragraph.add_run(node.title)

            # Add bookmark for document title
            if self.config.generate_bookmarks:
                self._add_bookmark_to_paragraph(paragraph, node.title)

        # Render children
        for child in node.children:
            self._render_node(child)

    def _render_section(self, node: SectionNode) -> None:
        """Render a section node."""
        assert self._document is not None
        assert self._resolver is not None

        number = self._resolver.section_numbers[id(node)]
        style = self.config.paragraph_styles.get_section_style(node.level)

        # Build section heading text
        if node.title:
            # Don't add extra period if number already has suffix
            if number.endswith(".") or number.endswith(")"):
                heading_text = f"{number} {node.title}"
            else:
                heading_text = f"{number}. {node.title}"
        else:
            # Section without title
            if number.endswith(".") or number.endswith(")"):
                heading_text = number
            else:
                heading_text = f"{number}."

        paragraph = self._document.add_paragraph(style=style)
        paragraph.add_run(heading_text)

        # Add bookmark for titled sections
        if node.title and self.config.generate_bookmarks:
            self._add_bookmark_to_paragraph(paragraph, node.title)

        # Render section content - collect inline children together
        inline_runs: list[Node] = []
        for child in node.children:
            if isinstance(
                child, TextNode | DefinedTermNode | CrossReferenceNode | CommentNode
            ):
                inline_runs.append(child)
            else:
                # Flush inline runs as a paragraph first
                if inline_runs:
                    self._render_inline_nodes_as_paragraph(inline_runs)
                    inline_runs = []
                # Render block-level node
                self._render_node(child)

        # Flush any remaining inline runs
        if inline_runs:
            self._render_inline_nodes_as_paragraph(inline_runs)

    def _render_paragraph(self, node: ParagraphNode) -> None:
        """Render a paragraph node."""
        assert self._document is not None

        style = self.config.paragraph_styles.paragraph
        paragraph = self._document.add_paragraph(style=style)
        for child in node.children:
            self._render_inline_node(child, paragraph)

    def _render_inline_nodes_as_paragraph(self, nodes: list[Node]) -> None:
        """Render a list of inline nodes as a single paragraph."""
        assert self._document is not None

        style = self.config.paragraph_styles.paragraph
        paragraph = self._document.add_paragraph(style=style)
        for node in nodes:
            self._render_inline_node(node, paragraph)

    def _render_inline_node(self, node: Node, paragraph: Paragraph) -> None:
        """Render an inline node into a paragraph."""
        if isinstance(node, TextNode):
            paragraph.add_run(node.text)
        elif isinstance(node, DefinedTermNode):
            self._render_defined_term(node, paragraph)
        elif isinstance(node, CrossReferenceNode):
            self._render_cross_reference(node, paragraph)
        elif isinstance(node, CommentNode) and self.config.include_comments:
            paragraph.add_run(f"[{node.content}]")

    def _render_defined_term(
        self, node: DefinedTermNode, paragraph: Paragraph
    ) -> None:
        """Render a defined term inline."""
        if node.descriptor:
            paragraph.add_run(f"({node.descriptor} ")
        else:
            paragraph.add_run("(")

        run = paragraph.add_run(node.term)
        if self.config.defined_term_bold:
            run.bold = True

        paragraph.add_run(")")

    def _render_cross_reference(
        self, node: CrossReferenceNode, paragraph: Paragraph
    ) -> None:
        """Render a cross-reference inline."""
        assert self._resolver is not None

        # Look up number
        number = self._resolver.title_to_number.get(node.reference_key)

        if number is None:
            # Missing reference - render original text
            paragraph.add_run(node.original_text)
            return

        # Strip trailing period from number for cross-references
        ref_number = number.rstrip(".") if number.endswith(".") else number

        # Apply template
        ref_text = self.config.cross_ref_template.format(number=ref_number)

        if self.config.generate_hyperlinks:
            anchor = self._resolver.title_to_anchor.get(node.reference_key)
            if anchor:
                bookmark_name = self._generate_bookmark_name(anchor)
                self._add_hyperlink(paragraph, ref_text, bookmark_name)
                return

        # No hyperlink, just add text
        paragraph.add_run(ref_text)

    def _render_comment(self, node: CommentNode) -> None:
        """Render a standalone comment node."""
        if not self.config.include_comments:
            return

        assert self._document is not None

        style = self.config.paragraph_styles.comment_text
        paragraph = self._document.add_paragraph(style=style)
        paragraph.add_run(node.content)

    def _render_signature_block(self, node: SignatureBlockNode) -> None:
        """Render a signature block."""
        assert self._document is not None

        styles = self.config.paragraph_styles

        if node.is_entity:
            # Entity signature
            party_name = node.party_name
            if self.config.uppercase_entity_names:
                party_name = party_name.upper()

            # Party name
            party_para = self._document.add_paragraph(style=styles.signature_party_name)
            run = party_para.add_run(party_name)
            run.bold = True

            # By Entity chain
            for by_entity in node.by_entities:
                entity_para = self._document.add_paragraph(
                    style=styles.signature_field
                )
                entity_para.add_run(f"By: {by_entity}")

            # Signature line
            sig_para = self._document.add_paragraph(style=styles.signature_line)
            sig_para.add_run("By: ")
            sig_line_run = sig_para.add_run("_" * 40)
            sig_line_run.underline = True

            # Human signatory
            if node.signatory:
                name_para = self._document.add_paragraph(style=styles.signature_field)
                name_para.add_run(f"Name: {node.signatory}")

            # Other fields
            for field_name, field_value in node.fields.items():
                if field_name.lower() not in ["by", "by entity"]:
                    field_para = self._document.add_paragraph(
                        style=styles.signature_field
                    )
                    field_para.add_run(f"{field_name}: {field_value}")

        else:
            # Individual signature
            # Signature line
            sig_para = self._document.add_paragraph(style=styles.signature_line)
            sig_line_run = sig_para.add_run("_" * 40)
            sig_line_run.underline = True

            # Party name
            name_para = self._document.add_paragraph(style=styles.signature_party_name)
            name_para.add_run(node.party_name)

            # Fields
            for field_name, field_value in node.fields.items():
                field_para = self._document.add_paragraph(style=styles.signature_field)
                field_para.add_run(f"{field_name}: {field_value}")

    def _add_bookmark_to_paragraph(self, paragraph: Paragraph, title: str) -> None:
        """Insert a bookmark at the start of a paragraph."""
        bookmark_name = self._generate_bookmark_name(title)

        # Create bookmark start element
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), str(id(paragraph)))
        bookmark_start.set(qn("w:name"), bookmark_name)

        # Create bookmark end element
        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), str(id(paragraph)))

        # Insert at beginning of paragraph
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)

    def _add_hyperlink(self, paragraph: Paragraph, text: str, anchor: str) -> None:
        """Add a hyperlink to an internal bookmark."""
        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)

        # Create run with text
        run = OxmlElement("w:r")

        # Add run properties for hyperlink styling
        run_props = OxmlElement("w:rPr")
        run_style = OxmlElement("w:rStyle")
        run_style.set(qn("w:val"), "Hyperlink")
        run_props.append(run_style)
        run.append(run_props)

        # Add text element
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        run.append(text_elem)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _generate_bookmark_name(self, title: str) -> str:
        """Generate a valid Word bookmark name from a title.

        Word bookmark constraints:
        - Max 40 characters
        - Must start with a letter
        - Can only contain letters, numbers, and underscores
        """
        # Normalize title: lowercase, replace spaces/special chars
        name = re.sub(r"[^\w\s]", "", title).strip()
        name = re.sub(r"[\s-]+", "_", name)
        name = name.lower()

        # Ensure starts with letter
        if not name or not name[0].isalpha():
            name = "ref_" + name

        # Truncate to 40 chars
        return name[:40]
